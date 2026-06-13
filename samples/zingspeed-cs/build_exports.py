#!/usr/bin/env python3
"""Dựng bộ tài liệu CS ZingSpeed ra nhiều định dạng — toàn bộ bằng Python.
Định dạng: .md (copy), .csv, .docx, .pdf, .png, .jpg, .jpeg
"""
import os, re, csv, shutil, html as htmllib

ROOT = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(ROOT, "exports")
IMG = os.path.join(OUT, "images")
MDD = os.path.join(OUT, "md")
for d in (OUT, IMG, MDD):
    os.makedirs(d, exist_ok=True)

FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"
FONT_IT = "/System/Library/Fonts/Supplemental/Arial Italic.ttf"
FONT_BI = "/System/Library/Fonts/Supplemental/Arial Bold Italic.ttf"
FONT_MONO = "/System/Library/Fonts/Supplemental/Courier New.ttf"
if not os.path.exists(FONT_MONO):
    FONT_MONO = FONT_REG

# Thứ tự ghép vào tài liệu tổng hợp
SOURCES = [
    ("README.md", "Tài liệu CS — ZingSpeed Mobile"),
    ("tai-lieu/00-tong-quan-game.md", "00 — Tổng quan game"),
    ("tai-lieu/01-tai-khoan-dang-nhap.md", "01 — Tài khoản & Đăng nhập"),
    ("tai-lieu/02-nap-the-giao-dich.md", "02 — Nạp thẻ & Giao dịch"),
    ("tai-lieu/03-loi-ky-thuat.md", "03 — Lỗi kỹ thuật"),
    ("tai-lieu/04-gameplay-co-ban.md", "04 — Gameplay cơ bản"),
    ("tai-lieu/05-su-kien-qua-tang.md", "05 — Sự kiện & Quà tặng"),
    ("tai-lieu/06-vi-pham-xu-ly.md", "06 — Vi phạm & Xử lý"),
    ("tai-lieu/07-kenh-ho-tro-escalation.md", "07 — Kênh hỗ trợ & Escalation"),
    ("NGUON-THAM-CHIEU.md", "Nguồn tham chiếu công khai"),
]
ALL_MD = [s[0] for s in SOURCES] + ["_TEMPLATE-chu-de.md"]

# ----------------------------------------------------------------------------
# 1) MARKDOWN -> BLOCK PARSER (tập con markdown dùng trong dự án)
# ----------------------------------------------------------------------------
def parse_blocks(md):
    lines = md.split("\n")
    blocks = []
    i = 0
    n = len(lines)
    in_comment = False
    while i < n:
        line = lines[i]
        # bỏ qua HTML comment <!-- ... -->
        if in_comment:
            if "-->" in line:
                in_comment = False
            i += 1
            continue
        if line.strip().startswith("<!--"):
            if "-->" not in line:
                in_comment = True
            i += 1
            continue
        s = line.strip()
        # fenced code
        if s.startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1
            blocks.append(("code", "\n".join(code)))
            continue
        # blank
        if s == "":
            i += 1; continue
        # hr
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", s):
            blocks.append(("hr", None)); i += 1; continue
        # heading
        m = re.match(r"(#{1,6})\s+(.*)", s)
        if m:
            blocks.append(("h", (len(m.group(1)), m.group(2).strip()))); i += 1; continue
        # table
        if "|" in s and i + 1 < n and re.fullmatch(r"\s*\|?[:\-\s|]+\|?\s*", lines[i+1]) and "-" in lines[i+1]:
            rows = []
            header = [c.strip() for c in s.strip().strip("|").split("|")]
            rows.append(header)
            i += 2  # bỏ dòng header + separator
            while i < n and "|" in lines[i] and lines[i].strip():
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells); i += 1
            blocks.append(("table", rows)); continue
        # blockquote
        if s.startswith(">"):
            qlines = []
            while i < n and lines[i].strip().startswith(">"):
                qlines.append(re.sub(r"^\s*>\s?", "", lines[i])); i += 1
            blocks.append(("quote", parse_blocks("\n".join(qlines)))); continue
        # unordered list
        if re.match(r"[-*]\s+", s):
            items = []
            while i < n and re.match(r"\s*[-*]\s+", lines[i]):
                items.append(re.sub(r"^\s*[-*]\s+", "", lines[i].strip())); i += 1
            blocks.append(("ul", items)); continue
        # ordered list
        if re.match(r"\d+\.\s+", s):
            items = []
            while i < n and re.match(r"\s*\d+\.\s+", lines[i]):
                items.append(re.sub(r"^\s*\d+\.\s+", "", lines[i].strip())); i += 1
            blocks.append(("ol", items)); continue
        # paragraph (gộp dòng liên tiếp)
        para = [s]
        i += 1
        while i < n and lines[i].strip() and not re.match(r"(#{1,6})\s|[-*]\s|\d+\.\s|>", lines[i].strip()) \
                and not lines[i].strip().startswith("```") and "|" not in lines[i]:
            para.append(lines[i].strip()); i += 1
        blocks.append(("p", " ".join(para)))
    return blocks

# ---- inline helpers ----
def _delink(s):
    s = re.sub(r"`([^`]*)`", r"\1", s)
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", s)
    return s

def segments(s):
    """-> list[(text, bold)] sau khi xử lý link/code/bold."""
    s = _delink(s)
    out = []
    bold = False
    for part in re.split(r"(\*\*)", s):
        if part == "**":
            bold = not bold; continue
        if part:
            out.append((part, bold))
    return out

def plain(s):
    return "".join(t for t, _ in segments(s))

def sanitize(s):
    """Thay emoji/ký hiệu Arial không có bằng glyph có sẵn — chỉ dùng cho PDF/ảnh."""
    s = s.replace("⚠️", "▲").replace("⚠", "▲")
    s = s.replace("✅", "●")
    s = s.replace("☐", "□")
    s = s.replace("‑", "-")   # non-breaking hyphen
    s = s.replace("️", "")    # variation selector
    return s

# ----------------------------------------------------------------------------
# 2) DOCX
# ----------------------------------------------------------------------------
def build_docx(path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def add_runs(par, text):
        for t, b in segments(text):
            r = par.add_run(t)
            r.bold = b

    first = True
    for relpath, title in SOURCES:
        if not first:
            doc.add_page_break()
        first = False
        md = open(os.path.join(ROOT, relpath), encoding="utf-8").read()
        for kind, payload in parse_blocks(md):
            if kind == "h":
                lvl, text = payload
                h = doc.add_heading(plain(text), level=min(lvl, 4))
            elif kind == "p":
                add_runs(doc.add_paragraph(), payload)
            elif kind == "ul":
                for it in payload:
                    add_runs(doc.add_paragraph(style="List Bullet"), it)
            elif kind == "ol":
                for it in payload:
                    add_runs(doc.add_paragraph(style="List Number"), it)
            elif kind == "quote":
                for k2, p2 in payload:
                    if k2 == "p":
                        par = doc.add_paragraph(style="Intense Quote")
                        add_runs(par, p2)
                    elif k2 in ("ul", "ol"):
                        for it in p2:
                            par = doc.add_paragraph(style="List Bullet" if k2 == "ul" else "List Number")
                            par.paragraph_format.left_indent = Inches(0.5)
                            add_runs(par, it)
                    elif k2 == "h":
                        add_runs(doc.add_paragraph(), p2[1])
            elif kind == "table":
                rows = payload
                t = doc.add_table(rows=0, cols=len(rows[0]))
                t.style = "Light Grid Accent 1"
                for ri, row in enumerate(rows):
                    cells = t.add_row().cells
                    for ci, cell in enumerate(row[:len(rows[0])]):
                        cells[ci].text = ""
                        add_runs(cells[ci].paragraphs[0], cell)
                        if ri == 0:
                            for run in cells[ci].paragraphs[0].runs:
                                run.bold = True
            elif kind == "code":
                par = doc.add_paragraph()
                run = par.add_run(payload)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            elif kind == "hr":
                doc.add_paragraph("─" * 40)
    doc.save(path)

# ----------------------------------------------------------------------------
# 3) PDF (fpdf2, font Arial Unicode tiếng Việt, inline bold qua markdown=True)
# ----------------------------------------------------------------------------
def build_pdf(path):
    from fpdf import FPDF
    pdf = FPDF(format="A4", unit="mm")
    pdf.set_auto_page_break(True, margin=15)
    pdf.add_font("Arial", "", FONT_REG)
    pdf.add_font("Arial", "B", FONT_BOLD)
    pdf.add_font("Arial", "I", FONT_IT)
    pdf.add_font("Arial", "BI", FONT_BI)
    pdf.add_font("Mono", "", FONT_MONO)
    EPW = 210 - 30  # effective page width

    def md_inline(s):
        # giữ ** cho bold, bỏ code backticks, link -> text (url)
        return _delink(s)

    first = True
    for relpath, title in SOURCES:
        pdf.add_page()
        first = False
        md = sanitize(open(os.path.join(ROOT, relpath), encoding="utf-8").read())
        for kind, payload in parse_blocks(md):
            if kind == "h":
                lvl, text = payload
                size = {1: 18, 2: 15, 3: 13}.get(lvl, 12)
                pdf.ln(2)
                pdf.set_font("Arial", "B", size)
                pdf.multi_cell(EPW, size * 0.45 + 2, plain(text))
                pdf.ln(1)
            elif kind == "p":
                pdf.set_font("Arial", "", 11)
                pdf.multi_cell(EPW, 6, md_inline(payload), markdown=True)
                pdf.ln(1)
            elif kind == "ul":
                pdf.set_font("Arial", "", 11)
                for it in payload:
                    pdf.multi_cell(EPW, 6, "•  " + md_inline(it), markdown=True)
                pdf.ln(1)
            elif kind == "ol":
                pdf.set_font("Arial", "", 11)
                for idx, it in enumerate(payload, 1):
                    pdf.multi_cell(EPW, 6, f"{idx}.  " + md_inline(it), markdown=True)
                pdf.ln(1)
            elif kind == "quote":
                pdf.set_font("Arial", "I", 11)
                for k2, p2 in payload:
                    if k2 == "p":
                        pdf.set_x(pdf.l_margin + 4)
                        pdf.multi_cell(EPW - 4, 6, md_inline(p2), markdown=True)
                    elif k2 in ("ul", "ol"):
                        for j, it in enumerate(p2, 1):
                            pdf.set_x(pdf.l_margin + 4)
                            bullet = "•  " if k2 == "ul" else f"{j}.  "
                            pdf.multi_cell(EPW - 4, 6, bullet + md_inline(it), markdown=True)
                pdf.ln(1)
            elif kind == "table":
                rows = payload
                pdf.set_font("Arial", "", 8.5)
                ncol = len(rows[0])
                with pdf.table(width=EPW, col_widths=tuple([1] * ncol),
                               text_align="LEFT", first_row_as_headings=True) as table:
                    for row in rows:
                        r = table.row()
                        for cell in row[:ncol]:
                            r.cell(plain(cell))
                pdf.ln(2)
            elif kind == "code":
                pdf.set_font("Mono", "", 8.5)
                for cl in payload.split("\n"):
                    pdf.multi_cell(EPW, 5, cl if cl else " ")
                pdf.ln(1)
            elif kind == "hr":
                y = pdf.get_y() + 1
                pdf.set_draw_color(180, 180, 180)
                pdf.line(pdf.l_margin, y, 210 - pdf.r_margin, y)
                pdf.ln(3)
    pdf.output(path)

# ----------------------------------------------------------------------------
# 4) ẢNH PNG/JPG/JPEG (Pillow) — mỗi file nguồn -> 1 ảnh dài
# ----------------------------------------------------------------------------
def build_images():
    from PIL import Image, ImageDraw, ImageFont
    W = 1100
    PAD = 50
    CW = W - 2 * PAD  # content width
    BG = (255, 255, 255)
    FG = (28, 28, 30)

    def F(path, size):
        return ImageFont.truetype(path, size)

    fonts = {
        "p": (F(FONT_REG, 22), F(FONT_BOLD, 22)),
        "h1": (F(FONT_BOLD, 38), F(FONT_BOLD, 38)),
        "h2": (F(FONT_BOLD, 30), F(FONT_BOLD, 30)),
        "h3": (F(FONT_BOLD, 25), F(FONT_BOLD, 25)),
        "mono": (F(FONT_MONO, 18), F(FONT_MONO, 18)),
        "quote": (F(FONT_IT, 22), F(FONT_BI, 22)),
        "tbl": (F(FONT_REG, 15), F(FONT_BOLD, 15)),
    }

    def text_w(draw, s, font):
        return draw.textbbox((0, 0), s, font=font)[2]

    def wrap_segments(draw, segs, font_pair, maxw):
        """Trả về list dòng; mỗi dòng = list (word, bold). Ngắt cả từ dài (URL)."""
        lines = [[]]
        cur = 0
        space_w = text_w(draw, " ", font_pair[0])
        for text, bold in segs:
            for word in re.split(r"(\s+)", text):
                if word == "" or word.isspace():
                    continue
                f = font_pair[1] if bold else font_pair[0]
                ww = text_w(draw, word, f)
                if ww > maxw:  # từ dài hơn cột (URL...) -> ngắt theo ký tự
                    piece = ""
                    for ch in word:
                        if piece and text_w(draw, piece + ch, f) > maxw:
                            if cur > 0:
                                lines.append([]); cur = 0
                            lines[-1].append((piece, bold))
                            lines.append([]); cur = 0
                            piece = ch
                        else:
                            piece += ch
                    if piece:
                        if cur > 0:
                            lines.append([]); cur = 0
                        lines[-1].append((piece, bold))
                        cur = text_w(draw, piece, f)
                    continue
                add = ww + (space_w if cur > 0 else 0)
                if cur + add > maxw and cur > 0:
                    lines.append([]); cur = 0
                    add = ww
                lines[-1].append((word, bold))
                cur += add
        return lines

    # tính chiều cao trước, rồi vẽ — dùng 1 hàm layout chung qua 2 lượt
    def layout(blocks, draw, measure_only, img=None):
        y = PAD
        x0 = PAD
        line_gap = 8

        def draw_line(words, font_pair, x, y, color=FG):
            space_w = text_w(draw, " ", font_pair[0])
            cx = x
            h = 0
            for word, bold in words:
                f = font_pair[1] if bold else font_pair[0]
                if not measure_only:
                    draw.text((cx, y), word, font=f, fill=color)
                ww = text_w(draw, word, f)
                bbox = draw.textbbox((0, 0), word, font=f)
                h = max(h, bbox[3])
                cx += ww + space_w
            return h

        def emit(segs, font_pair, indent=0, color=FG, maxw=None):
            nonlocal y
            mw = (CW - indent) if maxw is None else maxw
            for ln in wrap_segments(draw, segs, font_pair, mw):
                if not ln:
                    y += font_pair[0].size; continue
                h = draw_line(ln, font_pair, x0 + indent, y, color)
                y += (h if h else font_pair[0].size) + line_gap

        for kind, payload in blocks:
            if kind == "h":
                lvl, text = payload
                fp = fonts.get("h%d" % min(lvl, 3), fonts["h3"])
                y += 10
                emit(segments(text), fp)
                if lvl <= 2 and not measure_only:
                    draw.line((PAD, y + 2, W - PAD, y + 2), fill=(210, 210, 214), width=2)
                y += 8
            elif kind == "p":
                emit(segments(payload), fonts["p"])
                y += 4
            elif kind == "ul":
                for it in payload:
                    emit([("•  ", False)] + segments(it), fonts["p"], indent=10)
            elif kind == "ol":
                for idx, it in enumerate(payload, 1):
                    emit([("%d.  " % idx, False)] + segments(it), fonts["p"], indent=10)
            elif kind == "quote":
                top = y
                for k2, p2 in payload:
                    if k2 == "p":
                        emit(segments(p2), fonts["quote"], indent=22, color=(70, 70, 75))
                    elif k2 in ("ul", "ol"):
                        for j, it in enumerate(p2, 1):
                            b = "•  " if k2 == "ul" else "%d.  " % j
                            emit([(b, False)] + segments(it), fonts["quote"], indent=32, color=(70, 70, 75))
                if not measure_only:
                    draw.rectangle((PAD, top - 4, PAD + 5, y - line_gap + 4), fill=(120, 150, 220))
                y += 4
            elif kind == "table":
                rows = payload
                ncol = len(rows[0])
                colw = CW // ncol
                for ri, row in enumerate(rows):
                    fp = (fonts["tbl"][1], fonts["tbl"][1]) if ri == 0 else fonts["tbl"]
                    # tính chiều cao hàng
                    cell_lines = []
                    rowh = 0
                    for ci in range(ncol):
                        cell = row[ci] if ci < len(row) else ""
                        wl = wrap_segments(draw, segments(cell), fp, colw - 16)
                        cell_lines.append(wl)
                        rowh = max(rowh, len(wl) * (fp[0].size + 6) + 10)
                    if not measure_only:
                        draw.rectangle((PAD, y, W - PAD, y + rowh),
                                       outline=(200, 200, 205), width=1,
                                       fill=(238, 241, 248) if ri == 0 else None)
                    for ci in range(ncol):
                        cx = PAD + ci * colw + 8
                        cy = y + 6
                        for ln in cell_lines[ci]:
                            if not measure_only:
                                draw_line(ln, fp, cx, cy)
                            cy += fp[0].size + 6
                    y += rowh
                y += 8
            elif kind == "code":
                fp = fonts["mono"]
                lines = payload.split("\n")
                top = y
                if not measure_only:
                    h = len(lines) * (fp[0].size + 6) + 12
                    draw.rectangle((PAD, y, W - PAD, y + h), fill=(244, 244, 246))
                y += 6
                for cl in lines:
                    if not measure_only:
                        draw.text((PAD + 8, y), cl, font=fp[0], fill=(60, 60, 65))
                    y += fp[0].size + 6
                y += 10
            elif kind == "hr":
                if not measure_only:
                    draw.line((PAD, y + 4, W - PAD, y + 4), fill=(200, 200, 205), width=1)
                y += 14
        return y + PAD

    tmp = Image.new("RGB", (W, 10), BG)
    measure_draw = ImageDraw.Draw(tmp)

    produced = []
    for relpath, title in SOURCES:
        md = sanitize(open(os.path.join(ROOT, relpath), encoding="utf-8").read())
        blocks = parse_blocks(md)
        height = int(layout(blocks, measure_draw, measure_only=True))
        height = max(height, 200)
        img = Image.new("RGB", (W, height), BG)
        draw = ImageDraw.Draw(img)
        layout(blocks, draw, measure_only=False, img=img)
        base = os.path.splitext(os.path.basename(relpath))[0]
        png = os.path.join(IMG, base + ".png")
        img.save(png, "PNG")
        img.save(os.path.join(IMG, base + ".jpg"), "JPEG", quality=90)
        img.save(os.path.join(IMG, base + ".jpeg"), "JPEG", quality=90)
        produced.append((base, height))
    return produced

# ----------------------------------------------------------------------------
# 5) CSV (bảng hỏi-đáp cho CS)
# ----------------------------------------------------------------------------
def build_csv(path):
    TOPIC_FILES = [s[0] for s in SOURCES if s[0].startswith("tai-lieu/")]
    rows = []
    for relpath in TOPIC_FILES:
        fname = os.path.basename(relpath)
        md = open(os.path.join(ROOT, relpath), encoding="utf-8").read()
        # tách theo mục ###
        parts = re.split(r"\n### ", md)
        for part in parts[1:]:
            lines = part.split("\n")
            question = lines[0].strip()
            body = "\n".join(lines[1:])

            def grab(label, nexts):
                m = re.search(re.escape(label) + r"(.*?)(?=" + "|".join(re.escape(x) for x in nexts) + r"|\Z)",
                              body, re.S)
                if not m:
                    return ""
                txt = m.group(1).strip()
                txt = re.sub(r"^\*\*", "", txt)
                # gom dòng, bỏ ký hiệu blockquote/list thừa
                out = []
                for l in txt.split("\n"):
                    l = re.sub(r"^\s*>\s?", "", l)
                    l = l.strip()
                    if l:
                        out.append(l)
                return _delink(" ".join(out)).replace("**", "")

            answer = grab("**Trả lời được phép nói với user:**",
                          ["**⚠️ Ghi chú", "**Hướng xử lý", "**Nguồn:**"])
            internal = grab("**⚠️ Ghi chú nội bộ / KHÔNG nói với user:**",
                            ["**Hướng xử lý", "**Nguồn:**"])
            esc = grab("**Hướng xử lý / Escalation:**", ["**Nguồn:**"])
            src = grab("**Nguồn:**", ["**Kiểm chứng:**"])
            ver = grab("**Kiểm chứng:**", ["\n### ", "\n---"])
            rows.append({
                "File": fname,
                "Câu hỏi user": question,
                "Trả lời được phép nói với user": answer,
                "Ghi chú nội bộ (không nói user)": internal,
                "Hướng xử lý / Escalation": esc,
                "Nguồn": src,
                "Kiểm chứng & độ tin cậy": ver,
            })
    cols = ["File", "Câu hỏi user", "Trả lời được phép nói với user",
            "Ghi chú nội bộ (không nói user)", "Hướng xử lý / Escalation",
            "Nguồn", "Kiểm chứng & độ tin cậy"]
    with open(path, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.DictWriter(f, fieldnames=cols)
        w.writeheader()
        for r in rows:
            w.writerow(r)
    return len(rows)

# ----------------------------------------------------------------------------
# 6) Copy MD
# ----------------------------------------------------------------------------
def copy_md():
    for rel in ALL_MD:
        src = os.path.join(ROOT, rel)
        if os.path.exists(src):
            shutil.copy(src, os.path.join(MDD, os.path.basename(rel)))

# ----------------------------------------------------------------------------
if __name__ == "__main__":
    base = os.path.join(OUT, "ZingSpeed-CS-Tai-lieu")
    results = {}
    try:
        copy_md(); results["md"] = "OK -> exports/md/"
    except Exception as e:
        results["md"] = "LỖI: %r" % e
    try:
        n = build_csv(base + ".csv"); results["csv"] = "OK (%d mục) -> %s.csv" % (n, os.path.basename(base))
    except Exception as e:
        results["csv"] = "LỖI: %r" % e
    try:
        build_docx(base + ".docx"); results["docx"] = "OK -> %s.docx" % os.path.basename(base)
    except Exception as e:
        results["docx"] = "LỖI: %r" % e
    try:
        build_pdf(base + ".pdf"); results["pdf"] = "OK -> %s.pdf" % os.path.basename(base)
    except Exception as e:
        results["pdf"] = "LỖI: %r" % e
    try:
        imgs = build_images(); results["images"] = "OK (%d ảnh x3 định dạng) -> exports/images/" % len(imgs)
    except Exception as e:
        import traceback; traceback.print_exc()
        results["images"] = "LỖI: %r" % e
    print("\n=== KẾT QUẢ ===")
    for k in ("md", "csv", "docx", "pdf", "images"):
        print("%-8s %s" % (k, results.get(k, "—")))
