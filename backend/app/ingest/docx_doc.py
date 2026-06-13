"""DOCX text extraction (python-docx)."""

from __future__ import annotations

import io


def extract_docx(content: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
